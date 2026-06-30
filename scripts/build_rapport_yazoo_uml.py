from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\seef7\OneDrive\Desktop\YaZoo")
OUT_DIR = ROOT / "docs" / "rapport_yazoo_uml"
IMG_DIR = OUT_DIR / "uml"
DOCX_PATH = OUT_DIR / "Rapport_YaZoo_Analyse_UML.docx"

PURPLE = "6D28D9"
DARK = "241338"
LAVENDER = "F3ECFF"
MID = "8B5CF6"
TEXT = "1F2937"
MUTED = "6B7280"
BORDER = "D8B4FE"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "arialbd.ttf" if bold else "arial.ttf"
    path = Path("C:/Windows/Fonts") / name
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    fnt: ImageFont.ImageFont,
    fill: str = "#1F2937",
    align: str = "center",
    line_gap: int = 5,
) -> None:
    x1, y1, x2, y2 = box
    max_chars = max(8, int((x2 - x1) / max(7, fnt.size * 0.52)))
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap(part, max_chars) or [""])
    line_h = fnt.size + line_gap
    total_h = line_h * len(lines)
    y = y1 + max(0, (y2 - y1 - total_h) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        w = bbox[2] - bbox[0]
        if align == "left":
            x = x1 + 10
        elif align == "right":
            x = x2 - w - 10
        else:
            x = x1 + (x2 - x1 - w) // 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#6D28D9", width: int = 3) -> None:
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx > 0 else -1
        points = [(ex, ey), (ex - sign * 14, ey - 7), (ex - sign * 14, ey + 7)]
    else:
        sign = 1 if dy > 0 else -1
        points = [(ex, ey), (ex - 7, ey - sign * 14), (ex + 7, ey - sign * 14)]
    draw.polygon(points, fill=fill)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = "#FFFFFF",
    outline: str = "#D8B4FE",
    title_fill: str = "#6D28D9",
) -> None:
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    draw_wrapped(draw, title, (x1 + 8, y1 + 8, x2 - 8, y1 + 48), font(22, True), fill=title_fill)
    if body:
        draw.line((x1 + 14, y1 + 58, x2 - 14, y1 + 58), fill=outline, width=2)
        draw_wrapped(draw, body, (x1 + 12, y1 + 62, x2 - 12, y2 - 8), font(16), fill="#1F2937")


def actor(draw: ImageDraw.ImageDraw, center: tuple[int, int], label: str) -> None:
    x, y = center
    draw.ellipse((x - 18, y - 55, x + 18, y - 19), outline="#6D28D9", width=4)
    draw.line((x, y - 19, x, y + 36), fill="#6D28D9", width=4)
    draw.line((x - 35, y + 4, x + 35, y + 4), fill="#6D28D9", width=4)
    draw.line((x, y + 36, x - 30, y + 82), fill="#6D28D9", width=4)
    draw.line((x, y + 36, x + 30, y + 82), fill="#6D28D9", width=4)
    draw_wrapped(draw, label, (x - 80, y + 90, x + 80, y + 135), font(18, True), fill="#1F2937")


def ellipse_use_case(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], label: str) -> None:
    draw.ellipse(xy, fill="#FFFFFF", outline="#A78BFA", width=3)
    draw_wrapped(draw, label, xy, font(17, True), fill="#3B0764")


def base_canvas(title: str, size: tuple[int, int] = (1600, 1000)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", size, "#FBF7FF")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((30, 30, size[0] - 30, size[1] - 30), radius=28, fill="#FFFFFF", outline="#E9D5FF", width=3)
    draw.rectangle((30, 30, size[0] - 30, 115), fill="#F3ECFF")
    draw.text((70, 55), title, font=font(34, True), fill="#4C1D95")
    draw.text((size[0] - 340, 66), "YaZoo - UML", font=font(18, True), fill="#7C3AED")
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    path = IMG_DIR / name
    img.save(path, quality=95)
    return path


def make_use_case() -> Path:
    img, draw = base_canvas("Diagramme de cas d'utilisation")
    actor(draw, (145, 365), "Visiteur")
    actor(draw, (145, 690), "Utilisateur")
    actor(draw, (1455, 520), "Administrateur")
    cases = [
        ((420, 160, 700, 250), "S'inscrire / Se connecter"),
        ((770, 160, 1050, 250), "Consulter feed et profils"),
        ((420, 300, 700, 390), "Publier post / story"),
        ((770, 300, 1050, 390), "Créer communauté"),
        ((420, 450, 700, 540), "Publier annonce animal"),
        ((770, 450, 1050, 540), "Publier produit"),
        ((420, 600, 700, 690), "Réserver animal / produit"),
        ((770, 600, 1050, 690), "Messagerie privée"),
        ((420, 760, 700, 850), "Notifications"),
        ((1110, 360, 1390, 450), "Modérer contenus"),
        ((1110, 520, 1390, 610), "Gérer commandes"),
        ((1110, 690, 1390, 780), "Gérer utilisateurs"),
    ]
    for xy, label in cases:
        ellipse_use_case(draw, xy, label)
    for end in [(420, 205), (770, 205)]:
        arrow(draw, (250, 365), end, "#A78BFA", 2)
    for end in [(420, 345), (770, 345), (420, 495), (770, 495), (420, 645), (770, 645), (420, 805)]:
        arrow(draw, (250, 690), end, "#6D28D9", 2)
    for end in [(1390, 405), (1390, 565), (1390, 735)]:
        arrow(draw, (1360, 520), end, "#4C1D95", 2)
    return save(img, "01_use_case.png")


def make_class_diagram() -> Path:
    img, draw = base_canvas("Diagramme de classes métier")
    classes = {
        "User": ((80, 160, 345, 345), "+ posts()\n+ followers()\n+ products()\n+ reservationsAsBuyer()"),
        "Post": ((445, 150, 700, 330), "content\nmedia_path\nvisibility\n+ comments()\n+ likes()"),
        "Comment": ((790, 150, 1045, 330), "body\nreaction\nparent_id\n+ replies()"),
        "Like": ((1145, 150, 1400, 330), "reaction\nlikeable_type\nlikeable_id"),
        "Community": ((80, 470, 345, 650), "name\ndescription\nis_private\n+ memberships()"),
        "Animal": ((445, 455, 700, 650), "name\ncategory\nprice\nlisting_status\n+ reservations()"),
        "Product": ((790, 455, 1045, 650), "name\ncategory\nstock\nprice\n+ reservations()"),
        "Reservation": ((1145, 455, 1400, 650), "quantity\nstatus\npayment_status\ntotal_price\n+ reservable()"),
        "Conversation": ((250, 760, 520, 925), "participant_one_id\nparticipant_two_id\n+ messages()"),
        "Message": ((650, 760, 920, 925), "body\nread_at\n+ sender()"),
    }
    for name, (xy, body) in classes.items():
        rounded_box(draw, xy, name, body)
    relations = [
        ((345, 245), (445, 245), "1..*", 84),
        ((700, 245), (790, 245), "1..*", 84),
        ((1045, 245), (1145, 245), "0..*", 84),
        ((212, 345), (212, 470), "1..*", 84),
        ((345, 555), (445, 555), "1..*", 84),
        ((700, 555), (790, 555), "1..*", 84),
        ((1045, 555), (1145, 555), "0..*", 84),
        ((345, 300), (1230, 455), "buyer / seller", 130),
        ((520, 842), (650, 842), "1..*", 84),
        ((212, 650), (385, 760), "participants", 120),
    ]
    for start, end, label, label_width in relations:
        arrow(draw, start, end, "#7C3AED", 3)
        mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
        half_width = label_width // 2
        draw.rounded_rectangle((mx - half_width, my - 18, mx + half_width, my + 18), radius=8, fill="#F3ECFF")
        draw_wrapped(draw, label, (mx - half_width + 4, my - 15, mx + half_width - 4, my + 15), font(13, True), fill="#4C1D95")
    return save(img, "02_class_diagram.png")


def make_sequence_auth() -> Path:
    img, draw = base_canvas("Diagramme de séquence - Authentification")
    lanes = [170, 450, 730, 1010, 1290]
    labels = ["Utilisateur", "React/Vite", "API Laravel", "Sanctum", "MySQL"]
    for x, label in zip(lanes, labels):
        rounded_box(draw, (x - 95, 150, x + 95, 220), label, "", fill="#F3ECFF")
        draw.line((x, 220, x, 910), fill="#C4B5FD", width=3)
    steps = [
        (170, 450, 275, "Saisit email/téléphone + mot de passe"),
        (450, 730, 350, "POST /api/auth/login"),
        (730, 1290, 425, "Recherche utilisateur"),
        (1290, 730, 500, "Utilisateur + hash"),
        (730, 1010, 575, "Création token/cookie"),
        (730, 450, 650, "200 JSON {user} + cookie sécurisé"),
        (450, 730, 725, "GET /api/auth/me"),
        (730, 450, 800, "Profil authentifié"),
    ]
    for sx, ex, y, label in steps:
        arrow(draw, (sx, y), (ex, y), "#6D28D9", 3)
        draw_wrapped(draw, label, (min(sx, ex) + 15, y - 35, max(sx, ex) - 15, y - 8), font(15), fill="#1F2937")
    return save(img, "03_sequence_auth.png")


def make_sequence_post() -> Path:
    img, draw = base_canvas("Diagramme de séquence - Publication avec média")
    lanes = [150, 420, 690, 960, 1230, 1460]
    labels = ["Utilisateur", "CreatePost", "API Posts", "Storage", "PostResource", "Feed/Profile"]
    for x, label in zip(lanes, labels):
        rounded_box(draw, (x - 92, 150, x + 92, 220), label, "", fill="#F3ECFF")
        draw.line((x, 220, x, 910), fill="#C4B5FD", width=3)
    steps = [
        (150, 420, 280, "Sélection image/vidéo + texte"),
        (420, 690, 350, "POST /api/posts FormData"),
        (690, 960, 430, "Stocke media_path"),
        (960, 690, 505, "URL persistante /storage ou /media"),
        (690, 1230, 580, "Normalise auteur, média, compteurs"),
        (1230, 420, 655, "201 Post JSON"),
        (420, 1460, 730, "Mise à jour liste locale"),
        (1460, 690, 805, "Likes, commentaires, réactions"),
    ]
    for sx, ex, y, label in steps:
        arrow(draw, (sx, y), (ex, y), "#6D28D9", 3)
        draw_wrapped(draw, label, (min(sx, ex) + 10, y - 35, max(sx, ex) - 10, y - 8), font(14), fill="#1F2937")
    return save(img, "04_sequence_post_media.png")


def make_activity_reservation() -> Path:
    img, draw = base_canvas("Diagramme d'activité - Réservation")
    nodes = [
        ((120, 170, 410, 255), "Choisir animal ou produit"),
        ((120, 330, 410, 415), "Saisir quantité, livraison et note"),
        ((120, 490, 410, 575), "Créer réservation"),
        ((530, 490, 820, 575), "Notifier vendeur"),
        ((940, 330, 1230, 415), "Vendeur approuve ou rejette"),
        ((940, 490, 1230, 575), "Mettre à jour statut livraison"),
        ((940, 650, 1230, 735), "Finaliser commande"),
        ((530, 805, 820, 890), "Historique + facture"),
    ]
    for xy, label in nodes:
        rounded_box(draw, xy, label, "", fill="#FFFFFF")
    decision = (560, 300, 790, 430)
    draw.polygon([(675, 300), (790, 365), (675, 430), (560, 365)], fill="#F3ECFF", outline="#A78BFA")
    draw_wrapped(draw, "Stock / disponibilité OK ?", decision, font(17, True), fill="#4C1D95")
    for start, end in [((265, 255), (265, 330)), ((265, 415), (265, 490)), ((410, 532), (530, 532)), ((675, 430), (675, 490)), ((820, 532), (940, 372)), ((1085, 415), (1085, 490)), ((1085, 575), (1085, 650)), ((940, 692), (820, 847))]:
        arrow(draw, start, end, "#6D28D9", 3)
    arrow(draw, (410, 372), (560, 365), "#6D28D9", 3)
    draw_wrapped(draw, "Non : erreur de validation", (520, 220, 830, 270), font(15, True), fill="#B91C1C")
    arrow(draw, (675, 300), (675, 255), "#B91C1C", 3)
    return save(img, "05_activity_reservation.png")


def make_deployment() -> Path:
    img, draw = base_canvas("Diagramme de déploiement Azure / Docker")
    components = [
        ((70, 170, 360, 320), "Navigateur utilisateur", "HTTPS\nReact SPA\nCookies sécurisés"),
        ((500, 150, 850, 335), "Azure App Service\nFrontend yazoo", "Container Nginx\n5eef/yazoo-frontend\nCSP + assets statiques"),
        ((990, 150, 1340, 335), "Azure App Service\nBackend yazoo-api", "Container Nginx + PHP-FPM\nLaravel 12 API\nSanctum + Policies"),
        ((990, 450, 1340, 620), "Base de données MySQL", "Users, posts, marketplace,\ncommunities, reservations,\nmessages, notifications"),
        ((500, 450, 850, 620), "Docker Hub", "Images versionnées\n5eef/yazoo-api\n5eef/yazoo-frontend"),
        ((70, 450, 360, 620), "GitHub", "Code source\nActions / scripts\nhistorique Git"),
        ((500, 735, 850, 900), "Stockage média", "storage public Laravel\nimages posts, profils,\nannonces, communautés"),
    ]
    for xy, title, body in components:
        rounded_box(draw, xy, title, body)
    links = [
        ((360, 245), (500, 245), "HTTPS"),
        ((850, 245), (990, 245), "API /api"),
        ((1165, 335), (1165, 450), "SQL"),
        ((675, 450), (675, 335), "pull image"),
        ((215, 450), (500, 535), "CI/CD"),
        ((1165, 335), (675, 735), "media URLs"),
    ]
    for start, end, label in links:
        arrow(draw, start, end, "#6D28D9", 3)
        mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
        draw.rounded_rectangle((mx - 55, my - 18, mx + 55, my + 18), radius=10, fill="#F3ECFF")
        draw_wrapped(draw, label, (mx - 50, my - 15, mx + 50, my + 15), font(13, True), fill="#4C1D95")
    return save(img, "06_deployment_azure.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = TEXT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LAVENDER)
        set_cell_text(table.rows[0].cells[i], header, True, PURPLE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(PURPLE if level == 1 else DARK)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = "Caption"
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(MUTED)
        run.font.italic = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.5))
    add_caption(doc, caption)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [
        ("Title", 24, PURPLE),
        ("Subtitle", 13, MUTED),
        ("Heading 1", 17, PURPLE),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11, PURPLE),
        ("Caption", 9, MUTED),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if "Heading" in name or name == "Title":
            style.font.bold = True
    footer = section.footer.paragraphs[0]
    footer.text = "YaZoo - Rapport de soutenance 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OFFICE DE LA FORMATION PROFESSIONNELLE ET DE LA PROMOTION DU TRAVAIL")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(DARK)
    p = doc.add_paragraph("OFPPT - Développement Digital | Option Full-stack")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT DE SOUTENANCE\nPROJET DE FIN DE FORMATION")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(PURPLE)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("YaZoo - Réseau social et marketplace animalière")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor.from_string(DARK)
    p = doc.add_paragraph("Analyse, conception UML, réalisation et déploiement cloud Azure")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    values = [
        ("Projet", "YaZoo"),
        ("Stack", "React/Vite, Laravel/Sanctum, MySQL, Docker, Azure App Service"),
        ("Chemin analysé", str(ROOT)),
        ("Année", "2025 / 2026"),
    ]
    for row, (k, v) in zip(table.rows, values):
        set_cell_shading(row.cells[0], LAVENDER)
        set_cell_text(row.cells[0], k, True, PURPLE)
        set_cell_text(row.cells[1], v)
    doc.add_paragraph()
    p = doc.add_paragraph("Présenté par : Youssef Boughioul et équipe projet YaZoo")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p = doc.add_paragraph("Encadrement pédagogique : filière Développement Digital")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def build_doc(diagrams: list[Path]) -> None:
    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    sections = [
        ("Remerciements", [
            "Nous remercions l’équipe pédagogique de la filière Développement Digital pour l’accompagnement apporté durant la conception, le développement, les tests et le déploiement de YaZoo.",
            "Ce rapport s’appuie sur l’état réel du projet local analysé dans le dossier YaZoo, incluant le backend Laravel, le frontend React, les fichiers Docker, les routes API, les modèles de données et les mécanismes de déploiement Azure.",
        ]),
        ("Résumé", [
            "YaZoo est une application web sociale et marketplace destinée aux propriétaires et passionnés d’animaux. Elle rassemble les publications sociales, stories, profils, communautés, messagerie privée, annonces d’animaux et produits, réservations, notifications et administration.",
            "Le projet repose sur une architecture séparée : une interface React/Vite consomme une API REST Laravel sécurisée par Sanctum. Les médias sont gérés côté backend, les données principales sont stockées dans MySQL, et le déploiement est industrialisé avec Docker, Docker Hub et Azure App Service.",
        ]),
        ("Abstract", [
            "YaZoo is a pet-oriented social platform and marketplace. It centralizes social posts, stories, profiles, communities, private messaging, animal and product listings, reservations, notifications and administration.",
            "The solution is built with a separated architecture: a React/Vite frontend communicates with a Laravel REST API secured by Sanctum. Data is stored in MySQL, media are served by the backend, and deployment is prepared with Docker, Docker Hub and Azure App Service.",
        ]),
    ]
    for heading, paragraphs in sections:
        add_heading(doc, heading)
        for text in paragraphs:
            doc.add_paragraph(text)

    add_heading(doc, "Table des matières")
    doc.add_paragraph("Cliquez droit puis « Mettre à jour le champ » dans Word pour actualiser la table des matières si nécessaire.")
    add_heading(doc, "Liste des figures")
    for idx, caption in enumerate([
        "Diagramme de cas d’utilisation YaZoo",
        "Diagramme de classes métier",
        "Diagramme de séquence - authentification",
        "Diagramme de séquence - publication avec média",
        "Diagramme d’activité - réservation",
        "Diagramme de déploiement Azure / Docker",
    ], start=1):
        doc.add_paragraph(f"Figure {idx} : {caption}", style="List Number")
    add_heading(doc, "Liste des tableaux")
    for idx, caption in enumerate([
        "Réponse fonctionnelle aux besoins",
        "Modules fonctionnels",
        "Technologies utilisées",
        "Tables principales",
        "Tests et validation",
        "Risques et solutions",
    ], start=1):
        doc.add_paragraph(f"Tableau {idx} : {caption}", style="List Number")

    add_heading(doc, "Introduction générale")
    doc.add_paragraph(
        "Les usages numériques autour des animaux sont souvent dispersés entre réseaux sociaux, messageries, groupes communautaires et petites annonces. YaZoo propose de regrouper ces usages dans une plateforme unique, adaptée aux propriétaires d’animaux et aux acteurs du marché animalier."
    )
    doc.add_paragraph(
        "Le rapport présente l’analyse du besoin, la conception UML, les choix techniques, l’architecture logicielle, la réalisation, la sécurité, les tests et le déploiement de l’application."
    )

    add_heading(doc, "Présentation du projet YaZoo")
    add_table(doc, ["Besoin", "Réponse YaZoo"], [
        ["Publier des moments et conseils", "Feed social, posts avec médias, stories, commentaires, likes et réactions emoji."],
        ["Vendre ou proposer des animaux/produits", "Marketplace animaux et produits avec images, prix, statut, filtres et réservations."],
        ["Créer des groupes", "Communautés publiques ou privées, demandes d’adhésion et administration de groupe."],
        ["Échanger en privé", "Conversations privées entre utilisateurs, messages, notifications et lecture."],
        ["Administrer la plateforme", "Tableaux de bord pour modération des contenus et suivi des commandes."],
    ], [4.2, 11.8])

    add_heading(doc, "Contexte et problématique")
    add_bullets(doc, [
        "Les publications, annonces et discussions animalières sont souvent séparées entre plusieurs applications.",
        "Les utilisateurs ont besoin d’un profil centralisé avec leurs publications, annonces, communautés et réservations.",
        "Les transactions doivent être suivies via un flux de réservation clair : demande, approbation, livraison, facture et historique.",
        "La solution doit rester accessible depuis navigateur, sécurisée par API et déployable sur le cloud.",
    ])

    add_heading(doc, "Objectifs du projet")
    add_bullets(doc, [
        "Créer une expérience sociale animalière complète : feed, stories, profil, commentaires, likes et réactions.",
        "Intégrer une marketplace pour animaux et produits avec gestion des médias et réservations.",
        "Permettre les communautés de type groupe avec confidentialité, membres et publications associées.",
        "Assurer l’authentification email/téléphone/Google, la protection CORS/Sanctum et les réponses JSON normalisées.",
        "Préparer un déploiement reproductible avec Docker, Docker Hub et Azure App Service.",
    ])

    add_heading(doc, "Cahier des charges")
    add_heading(doc, "Acteurs", 2)
    add_bullets(doc, [
        "Visiteur : consulte la page d’accueil, s’inscrit, se connecte ou utilise Google OAuth.",
        "Utilisateur authentifié : publie, commente, réserve, échange par message, rejoint des communautés et gère son profil.",
        "Administrateur : modère les contenus, consulte les commandes et gère les utilisateurs.",
    ])
    add_heading(doc, "Exigences fonctionnelles", 2)
    add_table(doc, ["Module", "Fonctions principales"], [
        ["Authentification", "Inscription email/téléphone, connexion, logout, profil courant, Google OAuth, cookie Sanctum."],
        ["Feed social", "Posts, médias, visibilité, likes, réactions, commentaires, réponses et stories."],
        ["Profil", "Photo, couverture, bio, ville, statistiques dynamiques, publications récentes, bouton message."],
        ["Marketplace", "CRUD animaux/produits, filtres, images, adoption/vente, stock et statut."],
        ["Réservations", "Demande, approbation, rejet, annulation, livraison, finalisation, facture et historique."],
        ["Communautés", "Groupes publics/privés, demandes d’adhésion, membres, posts de groupe."],
        ["Messagerie", "Conversations privées, messages, notifications temps réel et ouverture depuis profil."],
    ], [3.7, 12.3])

    add_heading(doc, "Analyse fonctionnelle")
    add_figure(doc, diagrams[0], "Figure 1 : Diagramme de cas d’utilisation YaZoo.")
    doc.add_paragraph(
        "Le diagramme montre la séparation des responsabilités entre visiteur, utilisateur authentifié et administrateur. Les cas d’utilisation couvrent les modules réellement exposés par les routes API et les pages React."
    )

    add_heading(doc, "Analyse technique")
    add_table(doc, ["Couche", "Technologies et rôle"], [
        ["Frontend", "React, Vite, React Router, Axios, Tailwind CSS, Vitest, ESLint."],
        ["Backend", "Laravel 12, Sanctum, Socialite, Policies, Resources, Form Requests, Notifications."],
        ["Données", "MySQL pour les entités métier, migrations Laravel, relations Eloquent."],
        ["Temps réel", "Laravel Echo, Pusher, channels privés pour notifications et messages."],
        ["Médias", "Stockage public Laravel, URLs résolues par MediaStorage, CSP autorisant l’API Azure."],
        ["Déploiement", "Dockerfiles frontend/backend, Nginx, Docker Hub, Azure App Service."],
    ], [3.2, 12.8])

    add_heading(doc, "Architecture logicielle")
    doc.add_paragraph(
        "L’architecture est organisée autour d’une séparation nette entre interface et API. Le frontend centralise les appels HTTP dans des modules API, tandis que Laravel expose des routes REST protégées par Sanctum et des politiques d’autorisation."
    )
    add_bullets(doc, [
        "Le dossier frontend/src/pages contient les pages principales : feed, profil, marketplace, communautés, réservations, messagerie, paramètres et administration.",
        "Le dossier frontend/src/api regroupe les clients métier : auth, posts, stories, animals, products, reservations, communities, messages et notifications.",
        "Le backend sépare Controllers, Requests, Resources, Policies, Models, Notifications et Supports pour conserver une architecture maintenable.",
    ])

    add_heading(doc, "Architecture backend Laravel")
    doc.add_paragraph(
        "Le backend fournit 80 routes recensées par artisan route:list. Les routes principales sont préfixées par /api et couvrent l’authentification, les posts, les médias, les stories, les annonces, les réservations, les communautés, la messagerie et l’administration."
    )
    add_table(doc, ["Domaine API", "Endpoints représentatifs"], [
        ["Auth", "POST /api/auth/login, POST /api/auth/register, GET /api/auth/me, GET /api/auth/google."],
        ["Feed", "GET/POST /api/posts, PATCH/DELETE /api/posts/{post}, commentaires, likes, stories."],
        ["Marketplace", "GET/POST/PUT/DELETE /api/animals et /api/products."],
        ["Réservations", "POST /reservations, approve, reject, cancel, complete, delivery-status, invoice."],
        ["Communautés", "GET/POST/PUT /api/communities, join, leave, membership-requests."],
        ["Messages", "GET/POST /api/conversations, GET conversation, POST messages."],
    ], [4.0, 12.0])

    add_heading(doc, "Architecture frontend React")
    doc.add_paragraph(
        "Le frontend est une SPA structurée en pages, composants UI, hooks, contextes et modules API. Les pages utilisent React Router, les hooks centralisent l’authentification, la langue, le thème et les notifications, et Axios impose l’URL API configurée par VITE_API_URL."
    )
    add_bullets(doc, [
        "AuthContext : restauration utilisateur, CSRF, login/logout/register, synchronisation état local.",
        "ThemeContext et I18nContext : mode clair/sombre et langue choisie dans les paramètres.",
        "PostCard et CommentList : rendu des médias, likes, réactions emoji, commentaires et réponses.",
        "MarketplaceCards et formulaires : création d’annonces avec FormData lorsque des fichiers sont présents.",
    ])

    add_heading(doc, "Base de données")
    add_table(doc, ["Table / modèle", "Rôle"], [
        ["users", "Comptes, profil, admin, téléphone, langue, Google OAuth, avatar et couverture."],
        ["posts", "Publications sociales, médias, tags, visibilité et lien éventuel avec communauté."],
        ["comments / likes", "Commentaires, réponses, réactions emoji et likes polymorphes."],
        ["animals / products", "Annonces marketplace avec images, prix, statut, stock ou adoption."],
        ["reservations / reviews", "Commandes, livraison, paiement, facture et avis."],
        ["communities / community_members", "Groupes, confidentialité, rôles, statuts d’adhésion."],
        ["conversations / messages", "Messagerie privée entre deux participants."],
        ["notifications", "Notifications Laravel liées aux interactions métier."],
    ], [4.2, 11.8])

    add_heading(doc, "Diagramme de classes")
    add_figure(doc, diagrams[1], "Figure 2 : Diagramme de classes métier YaZoo.")
    doc.add_paragraph(
        "Le diagramme synthétise les relations Eloquent principales. Les réservations utilisent une relation polymorphe vers Animal ou Product. Les likes sont polymorphes et les commentaires supportent les réponses via parent_id."
    )

    add_heading(doc, "Séquence d’authentification")
    add_figure(doc, diagrams[2], "Figure 3 : Diagramme de séquence - authentification.")
    doc.add_paragraph(
        "L’authentification passe par l’API Laravel, qui valide les identifiants, interagit avec MySQL et renvoie un utilisateur JSON avec un cookie/token sécurisé compatible Sanctum."
    )

    add_heading(doc, "Séquence de publication avec média")
    add_figure(doc, diagrams[3], "Figure 4 : Diagramme de séquence - publication avec média.")
    doc.add_paragraph(
        "La création d’un post utilise FormData lorsque l’utilisateur joint une image ou une vidéo. Le backend stocke le fichier, persiste le chemin média, puis renvoie une ressource normalisée au feed et au profil."
    )

    add_heading(doc, "Activité de réservation")
    add_figure(doc, diagrams[4], "Figure 5 : Diagramme d’activité - réservation.")
    doc.add_paragraph(
        "La réservation couvre la disponibilité, la création de demande, la notification vendeur, l’approbation ou le rejet, le suivi livraison et l’émission de facture dans l’historique."
    )

    add_heading(doc, "Diagramme de déploiement")
    add_figure(doc, diagrams[5], "Figure 6 : Diagramme de déploiement Azure / Docker.")
    doc.add_paragraph(
        "Le déploiement de production repose sur deux App Services Azure : le frontend Nginx sert la SPA, tandis que le backend exécute Laravel avec Nginx et PHP-FPM. Les images sont versionnées dans Docker Hub."
    )

    add_heading(doc, "Réalisation du projet")
    add_heading(doc, "Interfaces principales", 2)
    add_bullets(doc, [
        "Landing page : présentation YaZoo, visuels animaliers, accès inscription/connexion.",
        "Feed : composer de posts/stories, liste dynamique, recherche, interactions et médias grand écran.",
        "Profil : couverture, avatar, statistiques followers/following/posts, publications récentes dynamiques et bouton de messagerie depuis un profil consulté.",
        "Marketplace : pages animaux et produits, formulaires de création, cartes, détails et réservations.",
        "Communautés : groupes disponibles, détail de communauté, confidentialité et publications de groupe.",
        "Paramètres : préférences utilisateur, thème et langue persistante.",
    ])
    add_heading(doc, "Sécurité", 2)
    add_bullets(doc, [
        "Sanctum protège les routes authentifiées et utilise cookies/tokens sécurisés.",
        "CORS autorise le frontend Azure et les environnements locaux configurés.",
        "SecurityHeaders et CSP réduisent les risques XSS, clickjacking et chargement de médias non autorisés.",
        "Policies Laravel contrôlent la modification/suppression des ressources appartenant aux utilisateurs.",
        "Les secrets ne doivent pas être versionnés ; les variables sensibles restent dans Azure App Settings ou .env local.",
    ])
    add_heading(doc, "Dockerisation et déploiement", 2)
    add_numbered(doc, [
        "Build de l’image backend depuis backend/Dockerfile avec Laravel, Nginx et PHP-FPM.",
        "Build de l’image frontend depuis frontend/Dockerfile avec Vite puis Nginx statique.",
        "Push des images vers Docker Hub sous le namespace 5eef.",
        "Configuration des App Services Azure yazoo et yazoo-api avec les tags Docker correspondants.",
        "Redémarrage, smoke tests HTTP et vérification CORS/API après déploiement.",
    ])

    add_heading(doc, "Tests et validation")
    add_table(doc, ["Type", "Commande", "Résultat observé"], [
        ["Backend", "php artisan test", "71 tests passés, 418 assertions."],
        ["Routes API", "php artisan route:list", "80 routes recensées, incluant auth, posts, communities, reservations, conversations."],
        ["Frontend lint", "npm run lint", "ESLint sans erreur bloquante."],
        ["Frontend tests", "npm test -- --run", "22 tests passés."],
        ["Frontend build", "npm run build", "Build Vite production généré avec succès."],
        ["Production", "curl Azure frontend/backend", "Réponses 200 et CORS actif côté API."],
    ], [3.2, 5.8, 7.0])

    add_heading(doc, "Difficultés rencontrées et solutions")
    add_table(doc, ["Difficulté", "Cause", "Solution appliquée"], [
        ["CORS bloquant les requêtes", "Origine frontend non autorisée ou credentials absents.", "Configuration CORS_ALLOWED_ORIGINS et supports_credentials."],
        ["Médias invisibles", "CSP frontend et URL médias API.", "CSP img-src/media-src autorisant yazoo-api et normalisation MediaStorage."],
        ["405 sur login/communautés", "Appels vers mauvais domaine ou routes sans /api.", "Client Axios centralisé vers VITE_API_URL."],
        ["Stats profil statiques", "Valeurs hardcodées.", "loadCount followers/following/posts et rendu dynamique."],
        ["Messagerie depuis profil", "Contact privé non exposé.", "Ouverture de conversation via recipient_id sécurisé."],
    ], [4.0, 5.5, 6.5])

    add_heading(doc, "Planification")
    add_table(doc, ["Phase", "Travaux réalisés"], [
        ["Analyse", "Identification des acteurs, modules, routes et modèles."],
        ["Conception", "UML, architecture API, schéma de données et flux métier."],
        ["Développement", "Frontend React, backend Laravel, formulaires, médias, notifications."],
        ["Tests", "PHPUnit, Vitest, lint, build, smoke tests production."],
        ["Déploiement", "Docker Hub, Azure App Service, variables production et redémarrage."],
        ["Amélioration", "Dark mode, i18n, Google OAuth, profil dynamique, communautés et messagerie."],
    ], [4.0, 12.0])

    add_heading(doc, "Bilan technique")
    doc.add_paragraph(
        "YaZoo présente une architecture suffisamment modulaire pour évoluer. Les modules fonctionnels sont isolés, l’API REST reste lisible, les données sont décrites par des migrations Laravel et les flux critiques sont couverts par des tests."
    )
    doc.add_paragraph(
        "Le projet montre également une capacité de déploiement réelle : Docker, Azure App Service, configuration CORS/CSP, stockage des médias et vérifications de production."
    )

    add_heading(doc, "Perspectives d’amélioration")
    add_bullets(doc, [
        "Ajouter un paiement en ligne avec un prestataire compatible Maroc.",
        "Migrer les médias vers Azure Blob Storage pour une scalabilité plus robuste.",
        "Renforcer la recherche avec indexation avancée et filtres géographiques.",
        "Ajouter une modération assistée et des signalements utilisateurs.",
        "Compléter l’i18n par une traduction exhaustive de tous les libellés.",
        "Mettre en place monitoring applicatif avec Application Insights.",
    ])

    add_heading(doc, "Conclusion générale")
    doc.add_paragraph(
        "Le projet YaZoo répond à un besoin concret : réunir dans une seule plateforme les usages sociaux, communautaires et commerciaux liés aux animaux. Le code analysé montre une solution full-stack fonctionnelle, structurée et déployable."
    )
    doc.add_paragraph(
        "Les six diagrammes UML produits dans ce rapport synthétisent la vision fonctionnelle, les classes métier, les flux d’authentification, de publication, de réservation et le déploiement Azure."
    )

    add_heading(doc, "Bibliographie et webographie")
    add_bullets(doc, [
        "Documentation Laravel : routing, Sanctum, Eloquent, Resources, Policies.",
        "Documentation React, Vite et React Router.",
        "Documentation Docker et Docker Compose.",
        "Documentation Microsoft Azure App Service.",
        "Code source local YaZoo analysé dans C:\\Users\\seef7\\OneDrive\\Desktop\\YaZoo.",
    ])

    add_heading(doc, "Annexes techniques")
    add_heading(doc, "Commandes utiles", 2)
    for cmd in [
        "cd backend && php artisan test",
        "cd backend && php artisan route:list",
        "cd frontend && npm run lint",
        "cd frontend && npm test -- --run",
        "cd frontend && npm run build",
        "docker build -f backend/Dockerfile -t 5eef/yazoo-api:<tag> .",
        "docker build -f frontend/Dockerfile -t 5eef/yazoo-frontend:<tag> .",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(cmd)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(DARK)

    add_heading(doc, "Glossaire", 2)
    add_table(doc, ["Terme", "Définition"], [
        ["SPA", "Application web monopage dont la navigation est gérée côté frontend."],
        ["API REST", "Interface HTTP structurée autour de ressources et méthodes standard."],
        ["Sanctum", "Solution Laravel pour authentification API par cookies ou tokens."],
        ["CSP", "Content Security Policy, politique limitant les sources de scripts, images et médias."],
        ["Docker", "Format de packaging permettant de déployer l’application avec ses dépendances."],
        ["App Service", "Service Azure d’hébergement d’applications web et containers."],
    ], [3.2, 12.8])

    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = [
        make_use_case(),
        make_class_diagram(),
        make_sequence_auth(),
        make_sequence_post(),
        make_activity_reservation(),
        make_deployment(),
    ]
    build_doc(diagrams)
    print(DOCX_PATH)
    for diagram in diagrams:
        print(diagram)


if __name__ == "__main__":
    main()
